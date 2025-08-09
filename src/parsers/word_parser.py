"""
Word文檔解析器
Word Document Parser

使用python-docx解析Word文檔，提取文本、圖片和表格。
嚴格按照項目規格文件的技術選型。

技術特點：
- 更好的格式支持和穩定性
- 完整的元數據提取
- 保留文檔結構信息
- 支持.docx和.doc格式
"""

from typing import List, Dict, Optional, Union, Any, Tuple
from pathlib import Path
import io
import hashlib
from dataclasses import dataclass

from src.parsers.base import (
    BaseParser, ParsedContent, TextBlock, ImageContent, TableContent,
    DocumentMetadata, ContentType, ImageFormat, ParserError,
    create_bounding_box, generate_content_id
)
from src.config.logging_config import get_logger, log_performance

logger = get_logger("word_parser")

class WordParser(BaseParser):
    """Word文檔解析器"""
    
    def __init__(self):
        """初始化Word解析器"""
        super().__init__("WordParser")
        
        self.supported_formats = ['.docx', '.doc']
        
        # 檢查依賴
        try:
            import docx
            self.docx = docx
            self.available = True
            logger.debug("python-docx可用")
        except ImportError:
            self.docx = None
            self.available = False
            logger.warning("python-docx不可用，請安裝: pip install python-docx")
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """
        檢查是否可以解析Word文件
        
        Args:
            file_path: 文件路徑
            
        Returns:
            bool: 是否可以解析
        """
        if not self.available:
            return False
        
        path = Path(file_path)
        return path.suffix.lower() in self.supported_formats and path.exists()
    
    @log_performance("word_parse")
    def parse(self, file_path: Union[str, Path], **kwargs) -> ParsedContent:
        """
        解析Word文檔
        
        Args:
            file_path: Word文件路徑
            **kwargs: 額外參數
            
        Returns:
            ParsedContent: 解析結果
        """
        if not self.available:
            raise ParserError("python-docx不可用", parser_name=self.name)
        
        self.validate_file(file_path)
        
        try:
            # 打開Word文檔
            doc = self.docx.Document(str(file_path))
            
            text_blocks = self._extract_text_blocks(doc)
            images = self._extract_images(doc)
            tables = self._extract_tables(doc)
            metadata = self._extract_metadata(doc, Path(file_path))
            
            result = ParsedContent(
                text_blocks=text_blocks,
                images=images,
                tables=tables,
                metadata=metadata,
                success=True
            )
            
            logger.info(f"Word解析完成: {len(text_blocks)}文本塊, {len(images)}圖片, {len(tables)}表格")
            return result
            
        except Exception as e:
            logger.error(f"Word解析失敗: {e}")
            return ParsedContent(
                success=False,
                error_message=f"Word解析失敗: {str(e)}"
            )
    
    def _extract_text_blocks(self, doc) -> List[TextBlock]:
        """提取文本塊"""
        text_blocks = []
        
        try:
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # 檢測內容類型
                content_type = ContentType.PARAGRAPH
                
                # 檢查是否為標題
                if paragraph.style.name.startswith('Heading'):
                    content_type = ContentType.HEADER
                    heading_level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                else:
                    heading_level = 0
                
                # 獲取格式信息
                font_name = ""
                font_size = 0.0
                is_bold = False
                is_italic = False
                
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    if first_run.font.name:
                        font_name = first_run.font.name
                    if first_run.font.size:
                        font_size = first_run.font.size.pt
                    is_bold = first_run.font.bold or False
                    is_italic = first_run.font.italic or False
                
                text_block = TextBlock(
                    id=generate_content_id("text", 1, i),
                    content=text,
                    content_type=content_type,
                    page_number=1,  # Word文檔沒有明確的頁面概念
                    paragraph_index=i,
                    font_name=font_name,
                    font_size=font_size,
                    is_bold=is_bold,
                    is_italic=is_italic,
                    heading_level=heading_level
                )
                
                text_blocks.append(text_block)
        
        except Exception as e:
            logger.warning(f"文本提取失敗: {e}")
        
        return text_blocks
    
    def _extract_images(self, doc) -> List[ImageContent]:
        """提取圖片"""
        images = []
        
        try:
            # 簡化的圖片提取（python-docx的圖片提取比較複雜）
            image_index = 0
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # 猜測圖片格式
                        if image_data.startswith(b'\xff\xd8'):
                            format_type = ImageFormat.JPEG
                            extension = "jpg"
                        elif image_data.startswith(b'\x89PNG'):
                            format_type = ImageFormat.PNG
                            extension = "png"
                        else:
                            format_type = ImageFormat.UNKNOWN
                            extension = "bin"
                        
                        image_content = ImageContent(
                            id=generate_content_id("image", 1, image_index),
                            filename=f"word_image_{image_index:03d}.{extension}",
                            format=format_type,
                            data=image_data,
                            page_number=1,
                            image_index=image_index
                        )
                        
                        images.append(image_content)
                        image_index += 1
                        
                    except Exception as e:
                        logger.warning(f"圖片 {image_index} 提取失敗: {e}")
                        continue
        
        except Exception as e:
            logger.warning(f"圖片提取失敗: {e}")
        
        return images
    
    def _extract_tables(self, doc) -> List[TableContent]:
        """提取表格"""
        tables = []
        
        try:
            for table_index, table in enumerate(doc.tables):
                try:
                    rows = []
                    
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        rows.append(row_data)
                    
                    if not rows:
                        continue
                    
                    # 檢測表頭
                    headers = []
                    table_rows = rows
                    
                    if len(rows) > 1:
                        # 假設第一行是表頭
                        potential_header = rows[0]
                        if all(cell and cell.strip() for cell in potential_header):
                            headers = potential_header
                            table_rows = rows[1:]
                    
                    table_content = TableContent(
                        id=generate_content_id("table", 1, table_index),
                        rows=table_rows,
                        headers=headers,
                        page_number=1,
                        table_index=table_index
                    )
                    
                    tables.append(table_content)
                    
                except Exception as e:
                    logger.warning(f"表格 {table_index} 提取失敗: {e}")
                    continue
        
        except Exception as e:
            logger.warning(f"表格提取失敗: {e}")
        
        return tables
    
    def _extract_metadata(self, doc, file_path: Path) -> DocumentMetadata:
        """提取文檔元數據"""
        try:
            core_props = doc.core_properties
            
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_path=str(file_path.absolute()),
                file_size=file_path.stat().st_size,
                file_format=file_path.suffix.lower(),
                title=core_props.title or "",
                author=core_props.author or "",
                subject=core_props.subject or "",
                creator=core_props.author or "",
                created_date=core_props.created,
                modified_date=core_props.modified,
                parser_name=self.name,
                parser_version="1.0.0"
            )
            
            # 處理關鍵詞
            if core_props.keywords:
                metadata.keywords = [kw.strip() for kw in core_props.keywords.split(",")]
            
            return metadata
            
        except Exception as e:
            logger.warning(f"元數據提取失敗: {e}")
            return DocumentMetadata(
                filename=file_path.name,
                file_path=str(file_path.absolute()),
                file_size=file_path.stat().st_size,
                file_format=file_path.suffix.lower(),
                parser_name=self.name,
                parser_version="1.0.0"
            )

def parse_word_document(file_path: Union[str, Path]) -> ParsedContent:
    """
    便捷的Word文檔解析函數
    
    Args:
        file_path: Word文件路徑
        
    Returns:
        ParsedContent: 解析結果
    """
    parser = WordParser()
    return parser.parse_with_metadata(file_path)

def extract_word_images(file_path: Union[str, Path]) -> List[ImageContent]:
    """
    提取Word文檔中的圖片
    
    Args:
        file_path: Word文件路徑
        
    Returns:
        List[ImageContent]: 圖片列表
    """
    parser = WordParser()
    result = parser.parse(file_path)
    return result.images if result.success else []

def extract_word_tables(file_path: Union[str, Path]) -> List[TableContent]:
    """
    提取Word文檔中的表格
    
    Args:
        file_path: Word文件路徑
        
    Returns:
        List[TableContent]: 表格列表
    """
    parser = WordParser()
    result = parser.parse(file_path)
    return result.tables if result.success else []

